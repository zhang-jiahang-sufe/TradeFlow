"""
报告导出工具 - 支持 Markdown、Word、PDF 格式

依赖安装:
    pip install pypandoc markdown

PDF 导出需要额外工具:
    - wkhtmltopdf (推荐): https://wkhtmltopdf.org/downloads.html
    - 或 LaTeX: https://www.latex-project.org/get/
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

# 检查依赖是否可用
try:
    import markdown
    import pypandoc

    # 检查 pandoc 是否可用
    try:
        pypandoc.get_pandoc_version()
        PANDOC_AVAILABLE = True
        logger.info("✅ Pandoc 可用")
    except OSError:
        PANDOC_AVAILABLE = False
        logger.warning("⚠️ Pandoc 不可用，Word 和 PDF 导出功能将不可用")

    EXPORT_AVAILABLE = True
except ImportError as e:
    EXPORT_AVAILABLE = False
    PANDOC_AVAILABLE = False
    logger.warning(f"⚠️ 导出功能依赖包缺失: {e}")
    logger.info("💡 请安装: pip install pypandoc markdown")

# 检查 pdfkit（唯一的 PDF 生成工具）
PDFKIT_AVAILABLE = False
PDFKIT_ERROR = None

try:
    import pdfkit
    # 检查 wkhtmltopdf 是否安装
    try:
        pdfkit.configuration()
        PDFKIT_AVAILABLE = True
        logger.info("✅ pdfkit + wkhtmltopdf 可用（PDF 生成工具）")
    except Exception as e:
        PDFKIT_ERROR = str(e)
        logger.warning("⚠️ wkhtmltopdf 未安装，PDF 导出功能不可用")
        logger.info("💡 安装方法: https://wkhtmltopdf.org/downloads.html")
except ImportError:
    logger.warning("⚠️ pdfkit 未安装，PDF 导出功能不可用")
    logger.info("💡 安装方法: pip install pdfkit")
except Exception as e:
    PDFKIT_ERROR = str(e)
    logger.warning(f"⚠️ pdfkit 检测失败: {e}")


class ReportExporter:
    """报告导出器 - 支持 Markdown、Word、PDF 格式"""

    def __init__(self):
        self.export_available = EXPORT_AVAILABLE
        self.pandoc_available = PANDOC_AVAILABLE
        self.pdfkit_available = PDFKIT_AVAILABLE

        logger.info("📋 ReportExporter 初始化:")
        logger.info(f"  - export_available: {self.export_available}")
        logger.info(f"  - pandoc_available: {self.pandoc_available}")
        logger.info(f"  - pdfkit_available: {self.pdfkit_available}")
    
    def generate_markdown_report(self, report_doc: Dict[str, Any]) -> str:
        """生成 Markdown 格式报告"""
        logger.info("📝 生成 Markdown 报告...")
        
        stock_symbol = report_doc.get("stock_symbol", "unknown")
        analysis_date = report_doc.get("analysis_date", "")
        analysts = report_doc.get("analysts", [])
        research_depth = report_doc.get("research_depth", 1)
        reports = report_doc.get("reports", {})
        summary = report_doc.get("summary", "")
        
        content_parts = []
        
        # 标题和元信息
        content_parts.append(f"# {stock_symbol} 股票分析报告")
        content_parts.append("")
        content_parts.append(f"**分析日期**: {analysis_date}")
        if analysts:
            content_parts.append(f"**分析师**: {', '.join(analysts)}")
        content_parts.append(f"**研究深度**: {research_depth}")
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        
        # 执行摘要
        if summary:
            content_parts.append("## 📊 执行摘要")
            content_parts.append("")
            content_parts.append(summary)
            content_parts.append("")
            content_parts.append("---")
            content_parts.append("")
        
        # 各模块内容
        module_order = [
            "company_overview",
            "financial_analysis", 
            "technical_analysis",
            "market_analysis",
            "risk_analysis",
            "valuation_analysis",
            "investment_recommendation"
        ]
        
        module_titles = {
            "company_overview": "🏢 公司概况",
            "financial_analysis": "💰 财务分析",
            "technical_analysis": "📈 技术分析",
            "market_analysis": "🌍 市场分析",
            "risk_analysis": "⚠️ 风险分析",
            "valuation_analysis": "💎 估值分析",
            "investment_recommendation": "🎯 投资建议"
        }
        
        # 按顺序添加模块
        for module_key in module_order:
            if module_key in reports:
                module_content = reports[module_key]
                if isinstance(module_content, str) and module_content.strip():
                    title = module_titles.get(module_key, module_key)
                    content_parts.append(f"## {title}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # 添加其他未列出的模块
        for module_key, module_content in reports.items():
            if module_key not in module_order:
                if isinstance(module_content, str) and module_content.strip():
                    content_parts.append(f"## {module_key}")
                    content_parts.append("")
                    content_parts.append(module_content)
                    content_parts.append("")
                    content_parts.append("---")
                    content_parts.append("")
        
        # 页脚
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        content_parts.append("*本报告由 TradingAgents-CN 自动生成*")
        content_parts.append("")
        
        markdown_content = "\n".join(content_parts)
        logger.info(f"✅ Markdown 报告生成完成，长度: {len(markdown_content)} 字符")
        
        return markdown_content
    
    def _clean_markdown_for_pandoc(self, md_content: str) -> str:
        """清理 Markdown 内容，避免 pandoc 解析问题"""
        import re

        # 移除可能导致 YAML 解析问题的内容
        # 如果开头有 "---"，在前面添加空行
        if md_content.strip().startswith("---"):
            md_content = "\n" + md_content

        # 🔥 移除可能导致竖排的 HTML 标签和样式
        # 移除 writing-mode 相关的样式
        md_content = re.sub(r'<[^>]*writing-mode[^>]*>', '', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<[^>]*text-orientation[^>]*>', '', md_content, flags=re.IGNORECASE)

        # 移除 <div> 标签中的 style 属性（可能包含竖排样式）
        md_content = re.sub(r'<div\s+style="[^"]*">', '<div>', md_content, flags=re.IGNORECASE)
        md_content = re.sub(r'<span\s+style="[^"]*">', '<span>', md_content, flags=re.IGNORECASE)

        # 🔥 移除可能导致问题的 HTML 标签
        # 保留基本的 Markdown 格式，移除复杂的 HTML
        md_content = re.sub(r'<style[^>]*>.*?</style>', '', md_content, flags=re.DOTALL | re.IGNORECASE)

        # 🔥 确保所有段落都是正常的横排文本
        # 在每个段落前后添加明确的换行，避免 Pandoc 误判
        lines = md_content.split('\n')
        cleaned_lines = []
        for line in lines:
            # 跳过空行
            if not line.strip():
                cleaned_lines.append(line)
                continue

            # 如果是标题、列表、表格等 Markdown 语法，保持原样
            if line.strip().startswith(('#', '-', '*', '|', '>', '```', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                cleaned_lines.append(line)
            else:
                # 普通段落：确保没有特殊字符导致竖排
                cleaned_lines.append(line)

        md_content = '\n'.join(cleaned_lines)

        return md_content

    def _create_pdf_css(self) -> str:
        """创建 PDF 样式表，控制表格分页和文本方向"""
        return """
<style>
/* 🔥 强制所有文本横排显示（修复中文竖排问题） */
* {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

body {
    writing-mode: horizontal-tb !important;
    direction: ltr !important;
}

/* 段落和文本 */
p, div, span, td, th, li {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表格样式 - 允许跨页 */
table {
    width: 100%;
    border-collapse: collapse;
    page-break-inside: auto;
    writing-mode: horizontal-tb !important;
}

/* 表格行 - 避免在行中间分页 */
tr {
    page-break-inside: avoid;
    page-break-after: auto;
}

/* 表头 - 在每页重复显示 */
thead {
    display: table-header-group;
}

/* 表格单元格 */
td, th {
    padding: 8px;
    border: 1px solid #ddd;
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
}

/* 表头样式 */
th {
    background-color: #f2f2f2;
    font-weight: bold;
}

/* 避免标题后立即分页 */
h1, h2, h3, h4, h5, h6 {
    page-break-after: avoid;
    writing-mode: horizontal-tb !important;
}

/* 避免在列表项中间分页 */
li {
    page-break-inside: avoid;
}

/* 代码块 */
pre, code {
    writing-mode: horizontal-tb !important;
    white-space: pre-wrap;
    word-wrap: break-word;
}
</style>
"""
    
    def generate_docx_report(self, report_doc: Dict[str, Any]) -> bytes:
        """生成 Word 文档格式报告"""
        logger.info("📄 开始生成 Word 文档...")

        if not self.pandoc_available:
            raise Exception("Pandoc 不可用，无法生成 Word 文档。请安装 pandoc 或使用 Markdown 格式导出。")

        # 生成 Markdown 内容
        md_content = self.generate_markdown_report(report_doc)

        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                output_file = tmp_file.name

            logger.info(f"📁 临时文件路径: {output_file}")

            # Pandoc 参数
            extra_args = [
                '--from=markdown-yaml_metadata_block',  # 禁用 YAML 元数据块解析
                '--standalone',  # 生成独立文档
                '--wrap=preserve',  # 保留换行
                '--columns=120',  # 设置列宽
                '-M', 'lang=zh-CN',  # 🔥 明确指定语言为简体中文
                '-M', 'dir=ltr',  # 🔥 明确指定文本方向为从左到右
            ]

            # 清理内容
            cleaned_content = self._clean_markdown_for_pandoc(md_content)

            # 转换为 Word
            pypandoc.convert_text(
                cleaned_content,
                'docx',
                format='markdown',
                outputfile=output_file,
                extra_args=extra_args
            )

            logger.info("✅ pypandoc 转换完成")

            # 🔥 后处理：修复 Word 文档中的文本方向
            try:
                from docx import Document
                doc = Document(output_file)

                # 修复所有段落的文本方向
                for paragraph in doc.paragraphs:
                    # 设置段落为从左到右
                    if paragraph._element.pPr is not None:
                        # 移除可能的竖排设置
                        for child in list(paragraph._element.pPr):
                            if 'textDirection' in child.tag or 'bidi' in child.tag:
                                paragraph._element.pPr.remove(child)

                # 修复表格中的文本方向
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if paragraph._element.pPr is not None:
                                    for child in list(paragraph._element.pPr):
                                        if 'textDirection' in child.tag or 'bidi' in child.tag:
                                            paragraph._element.pPr.remove(child)

                # 保存修复后的文档
                doc.save(output_file)
                logger.info("✅ Word 文档文本方向修复完成")
            except ImportError:
                logger.warning("⚠️ python-docx 未安装，跳过文本方向修复")
            except Exception as e:
                logger.warning(f"⚠️ Word 文档文本方向修复失败: {e}")

            # 读取生成的文件
            with open(output_file, 'rb') as f:
                docx_content = f.read()

            logger.info(f"✅ Word 文档生成成功，大小: {len(docx_content)} 字节")

            # 清理临时文件
            os.unlink(output_file)

            return docx_content
            
        except Exception as e:
            logger.error(f"❌ Word 文档生成失败: {e}", exc_info=True)
            # 清理临时文件
            try:
                if 'output_file' in locals() and os.path.exists(output_file):
                    os.unlink(output_file)
            except:
                pass
            raise Exception(f"生成 Word 文档失败: {e}")
    
    def _markdown_to_html(self, md_content: str) -> str:
        """将 Markdown 转换为 HTML"""
        import markdown

        # 配置 Markdown 扩展
        extensions = [
            'markdown.extensions.tables',  # 表格支持
            'markdown.extensions.fenced_code',  # 代码块支持
            'markdown.extensions.nl2br',  # 换行支持
        ]

        # 转换为 HTML
        html_content = markdown.markdown(md_content, extensions=extensions)

        # 添加 HTML 模板和样式
        # WeasyPrint 优化的 CSS（移除不支持的属性）
        html_template = f"""
<!DOCTYPE html>
<html lang="zh-CN" dir="ltr">
<head>
    <meta charset="UTF-8">
    <title>分析报告</title>
    <style>
        /* 基础样式 - 确保文本方向正确 */
        html {{
            direction: ltr;
        }}

        body {{
            font-family: "Noto Sans CJK SC", "Microsoft YaHei", "SimHei", "Arial", sans-serif;
            line-height: 1.8;
            color: #333;
            margin: 20mm;
            padding: 0;
            background: white;
            direction: ltr;
        }}

        /* 标题样式 */
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.8em;
            font-weight: 600;
            page-break-after: avoid;
            direction: ltr;
        }}

        h1 {{
            font-size: 2em;
            border-bottom: 3px solid #3498db;
            padding-bottom: 0.3em;
            page-break-before: always;
        }}

        h1:first-child {{
            page-break-before: avoid;
        }}

        h2 {{
            font-size: 1.6em;
            border-bottom: 2px solid #bdc3c7;
            padding-bottom: 0.25em;
        }}

        h3 {{
            font-size: 1.3em;
            color: #34495e;
        }}

        /* 段落样式 */
        p {{
            margin: 0.8em 0;
            text-align: left;
            direction: ltr;
        }}

        /* 表格样式 - 优化分页 */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1.5em 0;
            font-size: 0.9em;
            direction: ltr;
        }}

        /* 表头在每页重复 */
        thead {{
            display: table-header-group;
        }}

        tbody {{
            display: table-row-group;
        }}

        /* 表格行避免跨页断开 */
        tr {{
            page-break-inside: avoid;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 10px 12px;
            text-align: left;
            direction: ltr;
        }}

        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}

        tbody tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}

        tbody tr:hover {{
            background-color: #e9ecef;
        }}

        /* 代码块样式 */
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Consolas", "Monaco", "Courier New", monospace;
            font-size: 0.9em;
            direction: ltr;
        }}

        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #3498db;
            page-break-inside: avoid;
            direction: ltr;
        }}

        pre code {{
            background-color: transparent;
            padding: 0;
        }}

        /* 列表样式 */
        ul, ol {{
            margin: 0.8em 0;
            padding-left: 2em;
            direction: ltr;
        }}

        li {{
            margin: 0.4em 0;
            direction: ltr;
        }}

        /* 强调文本 */
        strong, b {{
            font-weight: 700;
            color: #2c3e50;
        }}

        em, i {{
            font-style: italic;
            color: #555;
        }}

        /* 水平线 */
        hr {{
            border: none;
            border-top: 2px solid #ecf0f1;
            margin: 2em 0;
        }}

        /* 链接样式 */
        a {{
            color: #3498db;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        /* 分页控制 */
        @page {{
            size: A4;
            margin: 20mm;

            @top-center {{
                content: "分析报告";
                font-size: 10pt;
                color: #999;
            }}

            @bottom-right {{
                content: "第 " counter(page) " 页";
                font-size: 10pt;
                color: #999;
            }}
        }}

        /* 避免孤行和寡行 */
        p, li {{
            orphans: 3;
            widows: 3;
        }}

        /* 图片样式 */
        img {{
            max-width: 100%;
            height: auto;
            page-break-inside: avoid;
        }}

        /* 引用块样式 */
        blockquote {{
            margin: 1em 0;
            padding: 0.5em 1em;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
            font-style: italic;
            page-break-inside: avoid;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
        return html_template

    def _generate_pdf_with_pdfkit(self, html_content: str) -> bytes:
        """使用 pdfkit 生成 PDF"""
        import pdfkit

        logger.info("🔧 使用 pdfkit + wkhtmltopdf 生成 PDF...")

        # 配置选项
        options = {
            'encoding': 'UTF-8',
            'enable-local-file-access': None,
            'page-size': 'A4',
            'margin-top': '20mm',
            'margin-right': '20mm',
            'margin-bottom': '20mm',
            'margin-left': '20mm',
        }

        # 生成 PDF
        pdf_bytes = pdfkit.from_string(html_content, False, options=options)

        logger.info(f"✅ pdfkit PDF 生成成功，大小: {len(pdf_bytes)} 字节")
        return pdf_bytes

    def generate_pdf_report(self, report_doc: Dict[str, Any]) -> bytes:
        """生成 PDF 格式报告（使用 pdfkit + wkhtmltopdf）"""
        logger.info("📊 开始生成 PDF 文档...")

        # 检查 pdfkit 是否可用
        if not self.pdfkit_available:
            error_msg = (
                "pdfkit 不可用，无法生成 PDF。\n\n"
                "安装方法:\n"
                "1. 安装 pdfkit: pip install pdfkit\n"
                "2. 安装 wkhtmltopdf: https://wkhtmltopdf.org/downloads.html\n"
            )
            if PDFKIT_ERROR:
                error_msg += f"\n错误详情: {PDFKIT_ERROR}"

            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)

        # 生成 Markdown 内容
        md_content = self.generate_markdown_report(report_doc)

        # 使用 pdfkit 生成 PDF
        try:
            html_content = self._markdown_to_html(md_content)
            return self._generate_pdf_with_pdfkit(html_content)
        except Exception as e:
            error_msg = f"PDF 生成失败: {e}"
            logger.error(f"❌ {error_msg}")
            raise Exception(error_msg)


# 创建全局导出器实例
report_exporter = ReportExporter()

